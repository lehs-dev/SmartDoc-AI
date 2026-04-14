import os
import pdfplumber
import docx
import ollama
import json
from collections import OrderedDict
from langchain_core.callbacks import StreamingStdOutCallbackHandler
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS
from langchain_ollama import OllamaLLM, OllamaEmbeddings
from langchain_core.prompts import PromptTemplate
from .models import ChatMessage, ChatSession, ConversationMemory, MemoryIndex

os.environ['HF_HUB_OFFLINE'] = '1'
VECTOR_DB_BASE_PATH = "vector_store"

SUPPORTED_LLM_MODELS = [
    "gemma4:e2b",
    "gemma4:e4b",
    "qwen3.5:2b",
    "qwen3.5:4b",
    "qwen3.5:9b",
]

SUPPORTED_EMBEDDING_MODELS = [
    "nomic-embed-text",      # ✅ Bạn đã pull
    "bge-m3:567m",           # ✅ Bạn đã pull
    "nomic-embed-text-v2-moe",
    "qwen3-embedding:0.6b",
]

VECTOR_DB_CONFIG = {
    "qwen_db": {
        "path": os.path.join(VECTOR_DB_BASE_PATH, "qwen_db"),
        "embedding_model": "qwen3-embedding:0.6b",
    },
    "bge_db": {
        "path": os.path.join(VECTOR_DB_BASE_PATH, "bge_db"),
        "embedding_model": "bge-m3:567m",
    },
    "nomic_v2_db": {
        "path": os.path.join(VECTOR_DB_BASE_PATH, "nomic_v2_db"),
        "embedding_model": "nomic-embed-text-v2-moe",
    },
    "nomic_v1_db": {
        "path": os.path.join(VECTOR_DB_BASE_PATH, "nomic_v1_db"),
        "embedding_model": "nomic-embed-text",
    },
}

_embedding_model_cache = {}
_llm_model_cache = {}
_vector_store_cache = {}
_installed_ollama_models_cache = None


def check_ollama_connection():
    """
    Kiểm tra kết nối Ollama
    Memory-Augmented RAG: Health check
    
    Returns:
        bool: True nếu Ollama đang chạy
    """
    try:
        print(f"\n🔍 [OLLAMA] Đang kiểm tra kết nối...")
        response = ollama.list()
        models = get_installed_ollama_models()
        print(f"✅ [OLLAMA] Kết nối thành công!")
        print(f"📦 Số models đã pull: {len(models)}")
        print(f"📋 Models: {', '.join(models[:5])}{'...' if len(models) > 5 else ''}")
        return True
    except Exception as e:
        print(f"❌ [OLLAMA] Không thể kết nối: {str(e)}")
        print(f"💡 Hint: Chạy 'ollama serve' hoặc kiểm tra Ollama đang chạy")
        return False
_memory_cache = OrderedDict()  # LRU cache cho memory
_memory_cache_max_size = 100  # Giới hạn số lượng memory items


def _extract_model_name(model_item):
    if isinstance(model_item, dict):
        return model_item.get("model") or model_item.get("name")

    return getattr(model_item, "model", None) or getattr(model_item, "name", None)


def _normalize_session_id(session_id):
    try:
        return int(session_id)
    except (TypeError, ValueError):
        return None


def _is_small_cpu_model(model_name):
    model_name = (model_name or '').lower()
    return 'e2b' in model_name or '0.8b' in model_name


def _truncate_text(text, max_chars):
    text = text or ''
    if len(text) <= max_chars:
        return text
    return text[:max_chars].rstrip() + '...'


def _format_recent_messages(messages, max_chars=1200):
    lines = []
    total_chars = 0

    for msg in messages:
        line = f"{msg.get_role_display()}: {msg.content}"
        if lines and total_chars + len(line) > max_chars:
            break
        lines.append(line)
        total_chars += len(line)

    return "\n".join(lines)


def _build_llm_kwargs(resolved_model_name):
    model_kwargs = {
        'model': resolved_model_name,
        'callbacks': [StreamingStdOutCallbackHandler()],
        'keep_alive': '15m',
    }

    if _is_small_cpu_model(resolved_model_name):
        model_kwargs.update({
            'temperature': 0.2,
            'repeat_penalty': 1.1,
            'top_k': 40,
            'top_p': 0.9,
            'num_ctx': 2048,
            'num_predict': 128,
            'num_thread': max(1, min(4, os.cpu_count() or 1)),
        })
    else:
        model_kwargs.update({
            'temperature': 0.7,
            'num_ctx': 4096,
            'num_thread': max(1, min(4, os.cpu_count() or 1)),
        })

    return model_kwargs


def get_installed_ollama_models(refresh=False):
    global _installed_ollama_models_cache

    if _installed_ollama_models_cache is not None and not refresh:
        return list(_installed_ollama_models_cache)

    try:
        response = ollama.list()
        model_items = []

        if isinstance(response, dict):
            model_items = response.get("models", [])
        else:
            model_items = getattr(response, "models", [])

        installed_models = []
        for item in model_items:
            model_name = _extract_model_name(item)
            if model_name:
                installed_models.append(model_name)

        _installed_ollama_models_cache = installed_models
    except Exception as e:
        print(f"Không lấy được danh sách model Ollama: {e}")
        _installed_ollama_models_cache = []

    return list(_installed_ollama_models_cache)


def _find_available_model(preferred_models):
    installed_models = get_installed_ollama_models()
    if not installed_models:
        return None

    installed_set = set(installed_models)
    for name in preferred_models:
        if name in installed_set:
            return name
    return None


def resolve_llm_model(model_name):
    _validate_llm_model(model_name)

    available = _find_available_model([model_name])
    if available:
        return available

    fallback = _find_available_model(SUPPORTED_LLM_MODELS)
    if fallback:
        print(f"Model {model_name} không có trong Ollama local, fallback sang {fallback}")
        return fallback

    raise ValueError(
        f"Model LLM '{model_name}' chưa có trong Ollama local. "
        "Vui lòng kiểm tra lại model đã pull về trước khi chat."
    )


def get_available_llm_models():
    installed_models = set(get_installed_ollama_models())
    if not installed_models:
        return list(SUPPORTED_LLM_MODELS)

    available = [model for model in SUPPORTED_LLM_MODELS if model in installed_models]
    return available or list(SUPPORTED_LLM_MODELS)


def get_available_embedding_models():
    return list(SUPPORTED_EMBEDDING_MODELS)


def _validate_llm_model(model_name):
    if model_name not in SUPPORTED_LLM_MODELS:
        raise ValueError(f"LLM model không được hỗ trợ: {model_name}")


def _validate_embedding_model(model_name):
    if model_name not in SUPPORTED_EMBEDDING_MODELS:
        raise ValueError(f"Embedding model không được hỗ trợ: {model_name}")


def resolve_vector_db_path(vector_db_key):
    config = VECTOR_DB_CONFIG.get(vector_db_key)
    if not config:
        raise ValueError(f"Vector DB key không hợp lệ: {vector_db_key}")
    return config["path"]


def route_embedding_target(file_size_bytes, has_vietnamese):
    file_size_mb = file_size_bytes / (1024 * 1024)

    # Quy tắc route theo yêu cầu hiện tại: tiếng Việt/file lớn ưu tiên Qwen,
    # tiếng Việt/file nhỏ ưu tiên BGE, tiếng Anh ưu tiên họ Nomic.
    if has_vietnamese:
        vector_db_key = "qwen_db" if file_size_mb > 5 else "bge_db"
    else:
        vector_db_key = "nomic_v2_db" if file_size_mb > 5 else "nomic_v1_db"

    config = VECTOR_DB_CONFIG[vector_db_key]
    return {
        "file_size_mb": round(file_size_mb, 2),
        "vector_db_key": vector_db_key,
        "vector_db_path": config["path"],
        "embedding_model": config["embedding_model"],
    }


def get_embeddings_model(model_name):
    """
    Load embedding model từ Ollama với caching
    Memory-Augmented RAG: Embedding model management
    
    Args:
        model_name: Tên embedding model
    
    Returns:
        OllamaEmbeddings instance
    """
    _validate_embedding_model(model_name)
    
    if model_name not in _embedding_model_cache:
        print(f"\n🤖 [EMBEDDING] Đang nạp model vào RAM: {model_name}...")
        try:
            _embedding_model_cache[model_name] = OllamaEmbeddings(model=model_name)
            print(f"✅ [EMBEDDING] Model loaded thành công: {model_name}")
        except Exception as e:
            print(f"❌ [EMBEDDING] Lỗi khi load model {model_name}: {str(e)}")
            import traceback
            traceback.print_exc()
            raise
    
    return _embedding_model_cache[model_name]


def get_llm_model(model_name):
    resolved_model_name = resolve_llm_model(model_name)

    if resolved_model_name not in _llm_model_cache:
        print(f"Khởi tạo kết nối tới Ollama với model: {resolved_model_name}...")
        _llm_model_cache[resolved_model_name] = OllamaLLM(**_build_llm_kwargs(resolved_model_name))
        
    return _llm_model_cache[resolved_model_name]


def get_cached_vector_store(vector_db_key, embedding_model_name):
    """Load FAISS theo từng kho và embedding model, chỉ 1 lần cho mỗi cặp."""
    cache_key = f"{vector_db_key}::{embedding_model_name}"
    if cache_key in _vector_store_cache:
        return _vector_store_cache[cache_key]

    print(f"Đang nạp Vector Database vào RAM: {vector_db_key}...")
    vector_db_path = resolve_vector_db_path(vector_db_key)
    index_path = os.path.join(vector_db_path, "index.faiss")

    if not os.path.exists(index_path):
        _vector_store_cache[cache_key] = None
        return None

    embeddings = get_embeddings_model(embedding_model_name)
    _vector_store_cache[cache_key] = FAISS.load_local(
        vector_db_path,
        embeddings,
        allow_dangerous_deserialization=True
    )
    return _vector_store_cache[cache_key]


def extract_text(file_path, file_extension):
    """
    Extract text từ PDF hoặc DOCX
    Memory-Augmented RAG: Text extraction với logging chi tiết
    
    Args:
        file_path: Đường dẫn file
        file_extension: 'pdf' hoặc 'docx'
    
    Returns:
        Extracted text
    """
    print(f"\n📄 [EXTRACT] Đang trích xuất text từ: {file_path}")
    print(f"📎 File extension: {file_extension}")
    
    text = ""
    try:
        if file_extension == 'pdf':
            print(f"📕 Processing PDF...")
            with pdfplumber.open(file_path) as pdf:
                print(f"📊 Số trang: {len(pdf.pages)}")
                for i, page in enumerate(pdf.pages):
                    extracted = page.extract_text()
                    if extracted:
                        text += extracted + "\n"
                        print(f"  ✅ Trang {i+1}: {len(extracted)} ký tự")
                    else:
                        print(f"  ⚠️  Trang {i+1}: Không extract được text")
                        
        elif file_extension == "docx":
            print(f"📘 Processing DOCX...")
            doc = docx.Document(file_path)
            print(f"📊 Số paragraphs: {len(doc.paragraphs)}")
            for i, para in enumerate(doc.paragraphs):
                if para.text.strip():
                    text += para.text + "\n"
                    if i < 5:  # Chỉ log 5 paragraphs đầu
                        print(f"  ✅ Paragraph {i+1}: {len(para.text)} ký tự")
        
        total_chars = len(text)
        print(f"\n✅ [EXTRACT] Hoàn thành! Tổng: {total_chars} ký tự")
        
        if total_chars == 0:
            print(f"⚠️  [EXTRACT] Cảnh báo: Không extract được text nào!")
        
        return text
        
    except Exception as e:
        print(f"❌ [EXTRACT] Lỗi khi đọc file {file_path}: {str(e)}")
        import traceback
        traceback.print_exc()
        return ""

def get_text_chunks(text):
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=1111,
        chunk_overlap=50,
        length_function=len,
    )
    chunks = text_splitter.split_text(text)
    return chunks


def get_text_chunks_optimized(text, file_size_mb=0, has_vietnamese=False):
    """
    Adaptive chunking dựa trên loại nội dung và kích thước file
    Tối ưu cho CPU/16GB RAM:
    - Giảm chunk size cho file lớn → giảm memory footprint
    - Tăng chunk overlap cho tiếng Việt → cải thiện context
    - Sử dụng separators phù hợp với ngôn ngữ
    
    Memory-Augmented RAG: Tối ưu chunking cho memory efficiency
    
    Args:
        text: Văn bản cần chunk
        file_size_mb: Kích thước file (MB)
        has_vietnamese: Có chứa tiếng Việt không
    
    Returns:
        list of text chunks
    """
    # Adaptive chunk size dựa trên file size và ngôn ngữ
    if file_size_mb > 10:
        # File lớn → chunk nhỏ để giảm RAM
        chunk_size = 512
        chunk_overlap = 30
    elif file_size_mb > 5:
        # File trung bình
        chunk_size = 768
        chunk_overlap = 40
    elif has_vietnamese:
        # Tiếng Việt cần nhiều context hơn
        chunk_size = 896
        chunk_overlap = 50
    else:
        # Mặc định
        chunk_size = 1024
        chunk_overlap = 50
    
    # Separators tối ưu cho tiếng Việt và English
    separators = [
        "\n\n",      # Paragraph breaks
        "\n",        # Line breaks
        "。",        # Chinese/Japanese period
        "!",        # Exclamation
        "?",        # Question
        ".",        # English period
        " ",        # Spaces
        ""          # Character level
    ]
    
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        length_function=len,
        separators=separators,
        is_separator_regex=False
    )
    
    chunks = text_splitter.split_text(text)
    
    print(f"Adaptive chunking: {len(chunks)} chunks | "
          f"size={chunk_size} | overlap={chunk_overlap} | "
          f"file_size={file_size_mb:.2f}MB | vietnamese={has_vietnamese}")
    
    return chunks

def process_document(file_path, file_extension):
    print(f'đang xử lý file: {file_path}...')
    raw_text = extract_text(file_path, file_extension)

    if not raw_text.strip():
        print(f"không tim thấy nội dung trong tài liệu")
        return []
    
    chunks = get_text_chunks(raw_text)
    print(f"Đã băm tài liệu thành {len(chunks)}  đoạn nhỏ")
    return chunks


def process_document_optimized(file_path, file_extension, file_size_mb=0, has_vietnamese=False):
    """
    Optimized document processing với adaptive chunking
    Memory-Augmented RAG: Xử lý tài liệu thông minh
    
    Args:
        file_path: Đường dẫn file
        file_extension: PDF hoặc DOCX
        file_size_mb: Kích thước file (MB)
        has_vietnamese: Có tiếng Việt không
    
    Returns:
        list of text chunks
    """
    print(f'Đang xử lý file (optimized): {file_path}...')
    raw_text = extract_text(file_path, file_extension)

    if not raw_text.strip():
        print(f"Không tìm thấy nội dung trong tài liệu")
        return []
    
    # Sử dụng adaptive chunking
    chunks = get_text_chunks_optimized(raw_text, file_size_mb, has_vietnamese)
    
    print(f"Đã xử lý tài liệu thành {len(chunks)} đoạn (optimized)")
    return chunks

def get_vector_store(chunks, embedding_model_name, vector_db_key):
    """
    Tạo và lưu vector store vào FAISS
    Memory-Augmented RAG: Vectorization với logging chi tiết
    
    Args:
        chunks: List of text chunks
        embedding_model_name: Tên embedding model
        vector_db_key: Key của vector database
    
    Returns:
        FAISS vector store
    """
    print(f"\n{'='*60}")
    print(f"🚀 BẮT ĐẦU VECTORIZE DOCUMENT")
    print(f"{'='*60}")
    print(f"📊 Số chunks: {len(chunks)}")
    print(f"🤖 Embedding model: {embedding_model_name}")
    print(f"💾 Vector DB key: {vector_db_key}")
    
    try:
        vector_db_path = resolve_vector_db_path(vector_db_key)
        print(f"📁 Vector DB path: {vector_db_path}")
        
        # Tạo thư mục nếu chưa tồn tại
        os.makedirs(vector_db_path, exist_ok=True)
        print(f"✅ Đã tạo thư mục: {vector_db_path}")
        
        # Load embedding model
        print(f"\n⏳ Đang load embedding model...")
        embeddings = get_embeddings_model(embedding_model_name)
        print(f"✅ Embedding model loaded thành công")
        
        index_path = os.path.join(vector_db_path, "index.faiss")
        print(f"📄 Index path: {index_path}")
        
        # Tạo hoặc update vector store
        if os.path.exists(index_path):
            print(f"📌 Index đã tồn tại, đang merge thêm chunks...")
            vector_store = FAISS.load_local(
                vector_db_path,
                embeddings,
                allow_dangerous_deserialization=True
            )
            vector_store.add_texts(chunks)
            print(f"✅ Đã merge {len(chunks)} chunks vào index cũ")
        else:
            print(f"📌 Index mới, đang tạo từ đầu...")
            vector_store = FAISS.from_texts(chunks, embedding=embeddings)
            print(f"✅ Đã tạo index mới với {len(chunks)} chunks")
        
        # Lưu vector store
        print(f"\n💾 Đang lưu vector store vào disk...")
        vector_store.save_local(vector_db_path)
        print(f"✅ Lưu thành công vào: {vector_db_path}")
        
        # Update cache
        cache_key = f"{vector_db_key}::{embedding_model_name}"
        _vector_store_cache[cache_key] = vector_store
        print(f"✅ Update cache: {cache_key}")
        
        print(f"\n{'='*60}")
        print(f"✅ VECTORIZE HOÀN THÀNH!")
        print(f"{'='*60}\n")
        
        return vector_store
        
    except Exception as e:
        print(f"\n{'='*60}")
        print(f"❌ LỖI VECTORIZE: {str(e)}")
        print(f"{'='*60}\n")
        import traceback
        traceback.print_exc()
        raise


def ask_gemma(
    question,
    chat_history="",
    llm_model_name="gemma4:e4b",
    embedding_model_name="",
    vector_db_key="",
):
    print('Đang tìm kiếm thông tin cho câu hỏi...')

    if not embedding_model_name or not vector_db_key:
        return iter([
            "Xin lỗi, tôi chưa xác định được kho dữ liệu cho tài liệu này. "
            "Vui lòng tải lại tài liệu hoặc chọn tài liệu khác."
        ])

    # Dùng DB đã cache trong RAM thay vì đọc ổ cứng
    vector_store = get_cached_vector_store(vector_db_key, embedding_model_name)

    if vector_store is None:
        return iter([
            "Xin lỗi, kho dữ liệu của tài liệu này chưa có nội dung. "
            "Vui lòng tải tài liệu lên trước."
        ])

    # GIẢM TẢI CHO CPU: Chỉ lấy 2 đoạn văn bản liên quan nhất (thay vì 3) để Prompt ngắn lại
    retriever = vector_store.as_retriever(search_kwargs={'k' : 2})
    relevant_doc = retriever.invoke(question)

    context  = "\n\n".join([doc.page_content for doc in relevant_doc])

    prompt_template =  """Bạn là trợ lý AI tên SmartDoc AI.
Lịch sử chat:
{chat_history}

Thông tin tài liệu:
{context}

Câu hỏi: {question}
Trả lời:"""

    prompt = PromptTemplate(template=prompt_template, input_variables=["chat_history" ,"context", "question"])

    print('Gemma 4 đang suy nghĩ...')
    llm = get_llm_model(llm_model_name)
    chain = prompt | llm

    return chain.stream({
        "chat_history": chat_history, 
        "context" : context, 
        "question" : question
    })


# ============================================================================
# GENERAL CHAT MODE - DIRECT LLM (NO RAG)
# ============================================================================

def ask_llm_direct(
    question,
    chat_history="",
    llm_model_name="gemma4:e4b",
):
    """
    General Chat Mode - Hỏi đáp trực tiếp với LLM không cần RAG
    Sử dụng khi không có tài liệu nào được upload
    
    Args:
        question: Câu hỏi của user
        chat_history: Lịch sử chat (optional)
        llm_model_name: Tên LLM model
    
    Returns:
        Stream generator từ LLM
    """
    print('\n💬 [GENERAL CHAT] Đang trả lời không cần RAG...')
    
    if _is_small_cpu_model(llm_model_name):
        prompt_template = """Bạn là SmartDoc AI.

Lịch sử chat:
{chat_history}

Câu hỏi:
{question}

Trả lời ngắn gọn, rõ ràng, bằng tiếng Việt:"""
    else:
        prompt_template = """Bạn là SmartDoc AI - một trợ lý AI hữu ích, thân thiện và thông minh.

Bạn có thể:
- Trả lời câu hỏi kiến thức chung
- Giải thích khái niệm, ý tưởng
- Hỗ trợ viết code, phân tích vấn đề
- Trò chuyện tự nhiên như người trợ lý

Lịch sử chat gần đây (nếu có):
{chat_history}

Câu hỏi của người dùng:
{question}

Hãy trả lời một cách tự nhiên, hữu ích và thân thiện. Sử dụng tiếng Việt trừ khi người dùng yêu cầu ngôn ngữ khác.

Trả lời:"""

    prompt = PromptTemplate(
        template=prompt_template,
        input_variables=["chat_history", "question"]
    )
    
    print(f'🤖 [GENERAL CHAT] Sử dụng model: {llm_model_name}')
    llm = get_llm_model(llm_model_name)
    chain = prompt | llm
    
    return chain.stream({
        "chat_history": _truncate_text(chat_history, 900 if _is_small_cpu_model(llm_model_name) else 2200),
        "question": question
    })


# ============================================================================
# MEMORY-AUGMENTED RAG FUNCTIONS
# ============================================================================

def get_memory_cache_max_size():
    """Trả về kích thước tối đa của memory cache"""
    return _memory_cache_max_size


def update_memory_cache(session_id, memory_data):
    """
    Cập nhật memory cache với LRU eviction
    Memory-Augmented RAG: Quản lý bộ nhớ đệm cho hội thoại
    """
    global _memory_cache
    
    # Nếu đã tồn tại, move to end (most recently used)
    if session_id in _memory_cache:
        _memory_cache.move_to_end(session_id)
    else:
        # Evict oldest nếu vượt quá giới hạn
        while len(_memory_cache) >= _memory_cache_max_size:
            _memory_cache.popitem(last=False)
    
    _memory_cache[session_id] = memory_data
    return memory_data


def get_recent_conversation_history(session_id, limit=5):
    """
    Lấy lịch sử hội thoại gần đây (Short-term Memory)
    Memory-Augmented RAG: Conversation Buffer Memory
    """
    session_id = _normalize_session_id(session_id)
    if session_id is None:
        return []

    messages = ChatMessage.objects.filter(
        session_id=session_id
    ).order_by('-created_at')[:limit]
    
    # Đảo ngược để theo thứ tự thời gian
    return list(reversed(messages))


def get_or_create_conversation_memory(session_id):
    """
    Lấy hoặc tạo ConversationMemory (Long-term Memory)
    Memory-Augmented RAG: Summary Memory
    """
    session_id = _normalize_session_id(session_id)
    if session_id is None:
        print(f"Lỗi khi lấy memory: session_id không hợp lệ ({session_id})")
        return None

    try:
        memory = ConversationMemory.objects.get(session_id=session_id)
        return memory
    except ConversationMemory.DoesNotExist:
        # Tạo memory mới nếu chưa tồn tại
        try:
            session = ChatSession.objects.get(id=session_id)
        except ChatSession.DoesNotExist:
            print(f"Lỗi khi tạo memory: session {session_id} không tồn tại")
            return None

        memory = ConversationMemory.objects.create(
            session=session,
            memory_type='summary'
        )
        return memory


def compress_conversation_to_summary(messages, llm_model_name="gemma4:e2b"):
    """
    Nén lịch sử hội thoại thành summary bằng LLM
    Memory-Augmented RAG: Memory Compression
    """
    if not messages:
        return ""
    
    # Format conversation
    conversation_text = "\n".join([
        f"{msg.role}: {msg.content}" for msg in messages
    ])
    
    prompt_template = """
Bạn là trợ lý AI. Hãy tóm tắt hội thoại sau thành các ý chính quan trọng.
Chỉ giữ lại: sự kiện, thông tin, con số, tên riêng, khái niệm quan trọng.
Bỏ qua: lời chào hỏi, câu hỏi lặp, thông tin không quan trọng.

Hội thoại:
{conversation}

Tóm tắt (tiếng Việt, ngắn gọn):
"""
    
    prompt = PromptTemplate(
        template=prompt_template,
        input_variables=["conversation"]
    )
    
    try:
        llm = get_llm_model(llm_model_name)
        chain = prompt | llm
        summary = chain.invoke({"conversation": conversation_text})
        return summary.strip()
    except Exception as e:
        print(f"Lỗi khi nén bộ nhớ: {e}")
        return ""


def extract_key_facts_from_conversation(messages, llm_model_name="gemma4:e2b"):
    """
    Trích xuất các sự kiện quan trọng từ hội thoại
    Memory-Augmented RAG: Entity Memory
    """
    if not messages:
        return {}
    
    conversation_text = "\n".join([
        f"{msg.role}: {msg.content}" for msg in messages
    ])
    
    prompt_template = """
Trích xuất thông tin quan trọng từ hội thoại sau.
Trả về JSON với các key:
- "entities": danh sách tên riêng, địa điểm, tổ chức
- "facts": danh sách sự kiện, thông tin quan trọng
- "numbers": danh sách con số, thống kê, ngày tháng

Hội thoại:
{conversation}

JSON:
"""
    
    prompt = PromptTemplate(
        template=prompt_template,
        input_variables=["conversation"]
    )
    
    try:
        llm = get_llm_model(llm_model_name)
        chain = prompt | llm
        result = chain.invoke({"conversation": conversation_text})
        
        # Parse JSON
        result_text = result.strip()
        # Remove markdown code blocks nếu có
        if result_text.startswith("```"):
            result_text = result_text.split("```")[1]
            if result_text.startswith("json"):
                result_text = result_text[4:]
        
        facts_dict = json.loads(result_text.strip())
        return facts_dict
    except Exception as e:
        print(f"Lỗi khi trích xuất sự kiện: {e}")
        return {"entities": [], "facts": [], "numbers": []}


def update_conversation_memory(session_id, force_update=False):
    """
    Cập nhật ConversationMemory từ lịch sử hội thoại
    Memory-Augmented RAG: Memory Update Strategy
    """
    session_id = _normalize_session_id(session_id)
    if session_id is None:
        print(f"Lỗi khi update memory: session_id không hợp lệ ({session_id})")
        return None

    # Lấy toàn bộ messages
    messages = ChatMessage.objects.filter(
        session_id=session_id
    ).order_by('created_at')
    
    if messages.count() == 0:
        return None
    
    # Chỉ update nếu có >= 3 messages hoặc force_update
    if messages.count() < 3 and not force_update:
        return get_or_create_conversation_memory(session_id)
    
    # Lấy memory hiện tại
    memory = get_or_create_conversation_memory(session_id)
    if memory is None:
        return None
    
    # Nén conversation thành summary
    summary = compress_conversation_to_summary(messages)
    if summary:
        memory.summary = summary
    
    # Trích xuất key facts
    facts = extract_key_facts_from_conversation(messages)
    if facts:
        memory.key_facts = json.dumps(facts, ensure_ascii=False)
    
    memory.save(update_fields=['summary', 'key_facts', 'last_updated'])
    
    # Update cache
    update_memory_cache(session_id, {
        'summary': memory.summary,
        'key_facts': memory.key_facts,
        'user_preferences': memory.user_preferences
    })
    
    return memory


def retrieve_with_memory_augmentation(
    question,
    session_id,
    vector_store,
    k_chunks=2,
    k_memories=1
):
    """
    Memory-Augmented Retrieval: Kết hợp retrieval từ nhiều nguồn
    1. Short-term: Last N messages (Conversation Buffer)
    2. Long-term: ConversationMemory summary (Summary Memory)
    3. Semantic: FAISS document chunks (Semantic Memory)
    
    Returns:
        dict: {
            'recent_messages': list of ChatMessage,
            'memory_context': str (summary),
            'key_facts': dict,
            'document_chunks': list of Document chunks,
            'combined_context': str (tất cả context)
        }
    """
    result = {
        'recent_messages': [],
        'memory_context': '',
        'key_facts': {},
        'document_chunks': [],
        'combined_context': ''
    }

    session_id = _normalize_session_id(session_id)
    if session_id is None:
        return result
    
    # 1. Get short-term memory (recent messages)
    recent_messages = get_recent_conversation_history(session_id, limit=5)
    result['recent_messages'] = recent_messages
    
    # 2. Get long-term memory (summary + facts)
    try:
        conv_memory = get_or_create_conversation_memory(session_id)
        if conv_memory is not None:
            result['memory_context'] = conv_memory.summary
        
            if conv_memory.key_facts:
                try:
                    result['key_facts'] = json.loads(conv_memory.key_facts)
                except:
                    result['key_facts'] = {}
    except Exception as e:
        print(f"Lỗi khi lấy memory: {e}")
    
    # 3. Get semantic memory (FAISS retrieval)
    if vector_store:
        retriever = vector_store.as_retriever(search_kwargs={'k': k_chunks})
        relevant_docs = retriever.invoke(question)
        result['document_chunks'] = relevant_docs
    
    # 4. Build combined context
    context_parts = []
    
    # Add memory context first (priority)
    if result['memory_context']:
        context_parts.append(f"=== TÓM TẮT HỘI THOẠI ===\n{result['memory_context']}")
    
    # Add key facts
    if result['key_facts']:
        facts_text = "=== SỰ KIỆN QUAN TRỌNG ===\n"
        if result['key_facts'].get('entities'):
            facts_text += f"Entities: {', '.join(result['key_facts']['entities'])}\n"
        if result['key_facts'].get('facts'):
            facts_text += f"Facts: {'; '.join(result['key_facts']['facts'])}\n"
        if result['key_facts'].get('numbers'):
            facts_text += f"Numbers: {', '.join(result['key_facts']['numbers'])}\n"
        context_parts.append(facts_text)
    
    # Add document chunks
    if result['document_chunks']:
        doc_context = "\n\n".join([doc.page_content for doc in result['document_chunks']])
        context_parts.append(f"=== THÔNG TIN TÀI LIỆU ===\n{doc_context}")
    
    result['combined_context'] = "\n\n".join(context_parts)
    
    return result


def ask_gemma_with_memory(
    question,
    session_id,
    llm_model_name="gemma4:e4b",
    embedding_model_name="",
    vector_db_key="",
    use_memory_augmentation=True,
    is_rag_mode=False
):
    """
    Memory-Augmented RAG: Hỏi đáp với memory augmentation
    Thay thế cho ask_gemma() truyền thống
    
    Args:
        question: Câu hỏi của user
        session_id: ID của chat session
        llm_model_name: Tên LLM model
        embedding_model_name: Tên embedding model
        vector_db_key: Key của vector DB
        use_memory_augmentation: Có sử dụng memory augmentation không
        is_rag_mode: True nếu đang ở chế độ RAG (có document), False nếu general chat
    
    Returns:
        Stream generator từ LLM
    """
    # Nếu không ở RAG mode, fallback về general chat
    if not is_rag_mode:
        print('\n🔹 [MODE] General Chat - Không dùng RAG')
        # Lấy chat history từ memory
        history_limit = 2 if _is_small_cpu_model(llm_model_name) else 5
        recent_messages = get_recent_conversation_history(session_id, limit=history_limit)
        chat_history = _format_recent_messages(
            recent_messages,
            max_chars=700 if _is_small_cpu_model(llm_model_name) else 1800,
        )
        return ask_llm_direct(
            question=question,
            chat_history=chat_history,
            llm_model_name=llm_model_name
        )
    
    print('Đang tìm kiếm thông tin cho câu hỏi (Memory-Augmented RAG)...')
    
    if not embedding_model_name or not vector_db_key:
        # Fallback về general chat nếu không có embedding info
        print('⚠️  [RAG] Không có embedding info, fallback về general chat')
        history_limit = 2 if _is_small_cpu_model(llm_model_name) else 5
        recent_messages = get_recent_conversation_history(session_id, limit=history_limit)
        chat_history = _format_recent_messages(
            recent_messages,
            max_chars=700 if _is_small_cpu_model(llm_model_name) else 1800,
        )
        return ask_llm_direct(
            question=question,
            chat_history=chat_history,
            llm_model_name=llm_model_name
        )
    
    # Load vector store
    vector_store = get_cached_vector_store(vector_db_key, embedding_model_name)
    
    if vector_store is None:
        # Fallback về general chat nếu không có vector store
        print('⚠️  [RAG] Không có vector store, fallback về general chat')
        history_limit = 2 if _is_small_cpu_model(llm_model_name) else 5
        recent_messages = get_recent_conversation_history(session_id, limit=history_limit)
        chat_history = _format_recent_messages(
            recent_messages,
            max_chars=700 if _is_small_cpu_model(llm_model_name) else 1800,
        )
        return ask_llm_direct(
            question=question,
            chat_history=chat_history,
            llm_model_name=llm_model_name
        )
    
    # Memory-Augmented Retrieval
    if use_memory_augmentation:
        k_chunks = 1 if _is_small_cpu_model(llm_model_name) else 2
        retrieval_result = retrieve_with_memory_augmentation(
            question=question,
            session_id=session_id,
            vector_store=vector_store,
            k_chunks=k_chunks,
            k_memories=1
        )
        
        # Format chat history từ recent messages
        recent_messages = retrieval_result['recent_messages']
        chat_history = _format_recent_messages(
            recent_messages,
            max_chars=700 if _is_small_cpu_model(llm_model_name) else 1800,
        )
        
        # Sử dụng combined context từ memory + documents
        context = _truncate_text(
            retrieval_result['combined_context'],
            900 if _is_small_cpu_model(llm_model_name) else 3500,
        )
        
        # Update memory sau khi retrieve (async, không block)
        try:
            update_conversation_memory(session_id)
        except Exception as e:
            print(f"Lỗi khi update memory: {e}")
    
    else:
        # Fallback: Không dùng memory (như ask_gemma cũ)
        retriever = vector_store.as_retriever(search_kwargs={'k': 1 if _is_small_cpu_model(llm_model_name) else 2})
        relevant_docs = retriever.invoke(question)
        context = _truncate_text(
            "\n\n".join([doc.page_content for doc in relevant_docs]),
            900 if _is_small_cpu_model(llm_model_name) else 3500,
        )
        chat_history = ""
    
    # Build prompt với memory context
    prompt_template = """Bạn là trợ lý AI tên SmartDoc AI, sử dụng Memory-Augmented RAG.

Lịch sử chat gần đây:
{chat_history}

Thông tin từ bộ nhớ và tài liệu:
{context}

Câu hỏi: {question}

Hãy trả lời dựa trên thông tin trên, kết hợp với ngữ cảnh từ lịch sử chat.
Nếu thông tin không có trong tài liệu, hãy nói rõ và đưa ra câu trả lời chung.

Trả lời (tiếng Việt):"""

    prompt = PromptTemplate(
        template=prompt_template,
        input_variables=["chat_history", "context", "question"]
    )
    
    print('Gemma 4 đang suy nghĩ (với memory context)...')
    llm = get_llm_model(llm_model_name)
    chain = prompt | llm
    
    return chain.stream({
        "chat_history": _truncate_text(chat_history, 700 if _is_small_cpu_model(llm_model_name) else 1800),
        "context": context,
        "question": question
    })